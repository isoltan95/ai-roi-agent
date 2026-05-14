import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from .models import ROIResult

# NPC color palette
NPC_MAROON = RGBColor(0x8B, 0x17, 0x2E)   # Qatar maroon
NPC_NAVY = RGBColor(0x0D, 0x1B, 0x3E)      # Dark navy
NPC_GOLD = RGBColor(0xC9, 0xA0, 0x2B)      # Qatar gold
LIGHT_GRAY = RGBColor(0xF5, 0xF5, 0xF5)
TEXT_DARK = RGBColor(0x1A, 0x1A, 0x2E)


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc):
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "8B172E")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _heading(doc, text: str, level: int = 1):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = NPC_MAROON
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = NPC_NAVY
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = NPC_NAVY
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    return para


def _body(doc, text: str):
    para = doc.add_paragraph(text)
    para.runs[0].font.size = Pt(11)
    para.runs[0].font.color.rgb = TEXT_DARK
    para.paragraph_format.space_after = Pt(6)
    return para


def _kv_table(doc, rows: list[tuple[str, str]]):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for i, (key, val) in enumerate(rows):
        key_cell = table.rows[i].cells[0]
        val_cell = table.rows[i].cells[1]
        key_cell.text = key
        val_cell.text = val
        key_run = key_cell.paragraphs[0].runs[0]
        key_run.bold = True
        key_run.font.size = Pt(10)
        _set_cell_bg(key_cell, "F0F0F0")
        val_cell.paragraphs[0].runs[0].font.size = Pt(10)
        # Column widths
        key_cell.width = Inches(2.5)
        val_cell.width = Inches(4.0)
    doc.add_paragraph()


def _financial_table(doc, result: ROIResult):
    fm = result.financial_metrics
    headers = ["Financial Metric", "Value", "Notes"]
    rows_data = [
        ("3-Year ROI", f"{fm.three_year_roi_percentage:,.1f}%", "Return on total 3-year investment"),
        ("Payback Period", f"{fm.payback_period_months:.1f} months", "Time to recover implementation cost"),
        ("3-Year NPV (8%)", f"QAR {fm.npv_3yr_qar:,.0f}", "Net present value at 8% discount rate"),
        ("Benefit-Cost Ratio", f"{fm.benefit_cost_ratio:.2f}x", "Total benefits vs. total costs"),
        ("Annual Net Benefit", f"QAR {fm.annual_net_benefit_qar:,.0f}", "After maintenance costs"),
        ("Total Annual Benefits", f"QAR {result.total_annual_benefits_qar:,.0f}", "Gross annual benefit"),
        ("Annual Maintenance", f"QAR {result.annual_maintenance_cost_qar:,.0f}", "Recurring operational cost"),
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=3)
    table.style = "Table Grid"

    # Header row
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        _set_cell_bg(cell, "8B172E")

    for i, (metric, value, note) in enumerate(rows_data):
        row = table.rows[i + 1]
        row.cells[0].text = metric
        row.cells[1].text = value
        row.cells[2].text = note
        row.cells[1].paragraphs[0].runs[0].bold = True
        if i % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, "FAF0F2")
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()


def _benefits_table(doc, result: ROIResult):
    rows_data = [
        ("Labor / FTE Savings", result.annual_labor_savings_qar),
        ("Error Reduction Savings", result.annual_error_savings_qar),
        ("Productivity Gains", result.annual_productivity_savings_qar),
        ("Other Savings", result.annual_other_savings_qar),
        ("TOTAL ANNUAL BENEFITS", result.total_annual_benefits_qar),
    ]

    table = doc.add_table(rows=1 + len(rows_data), cols=2)
    table.style = "Table Grid"

    for j, h in enumerate(["Benefit Category", "Annual Value (QAR)"]):
        cell = table.rows[0].cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        _set_cell_bg(cell, "0D1B3E")

    for i, (label, val) in enumerate(rows_data):
        row = table.rows[i + 1]
        row.cells[0].text = label
        row.cells[1].text = f"QAR {val:,.0f}"
        is_total = i == len(rows_data) - 1
        for cell in row.cells:
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            if is_total:
                run.bold = True
                _set_cell_bg(cell, "FFF3CD")
    doc.add_paragraph()


def _risk_table(doc, risks: list):
    headers = ["Risk", "Description", "Likelihood", "Impact", "Mitigation"]
    table = doc.add_table(rows=1 + len(risks), cols=5)
    table.style = "Table Grid"

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        _set_cell_bg(cell, "8B172E")

    color_map = {"High": "FFCCCC", "Medium": "FFF3CD", "Low": "D4EDDA"}
    for i, risk in enumerate(risks):
        row = table.rows[i + 1]
        row.cells[0].text = risk.get("title", "")
        row.cells[1].text = risk.get("description", "")
        row.cells[2].text = risk.get("likelihood", "")
        row.cells[3].text = risk.get("impact", "")
        row.cells[4].text = risk.get("mitigation", "")
        lkl = risk.get("likelihood", "Low")
        _set_cell_bg(row.cells[2], color_map.get(lkl, "FFFFFF"))
        _set_cell_bg(row.cells[3], color_map.get(risk.get("impact", "Low"), "FFFFFF"))
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()


def generate_report(result: ROIResult) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── COVER PAGE ──────────────────────────────────────────────
    # Title block
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("مجلس التخطيط الوطني\nNATIONAL PLANNING COUNCIL")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NPC_MAROON

    sub = doc.add_paragraph("State of Qatar  |  دولة قطر")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = NPC_GOLD

    doc.add_paragraph()
    _add_horizontal_rule(doc)
    doc.add_paragraph()

    main_title = doc.add_paragraph("AI Use Case ROI Analysis Report")
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    main_title.runs[0].bold = True
    main_title.runs[0].font.size = Pt(22)
    main_title.runs[0].font.color.rgb = NPC_NAVY

    inp = result.use_case_input
    uc_title = doc.add_paragraph(inp.title)
    uc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    uc_title.runs[0].font.size = Pt(16)
    uc_title.runs[0].font.color.rgb = NPC_MAROON
    uc_title.runs[0].bold = True

    doc.add_paragraph()
    _add_horizontal_rule(doc)
    doc.add_paragraph()

    meta_rows = [
        ("Sector", inp.sector),
        ("Department", inp.department),
        ("Submitted by", inp.submitter_name),
        ("Date", datetime.now().strftime("%d %B %Y")),
        ("Classification", "CONFIDENTIAL – NPC INTERNAL USE ONLY"),
    ]
    _kv_table(doc, meta_rows)

    doc.add_page_break()

    # ── 1. EXECUTIVE SUMMARY ──────────────────────────────────
    _heading(doc, "1. Executive Summary")
    _add_horizontal_rule(doc)
    _body(doc, result.executive_summary)

    doc.add_paragraph()
    _heading(doc, "Key ROI Metrics at a Glance", level=2)
    fm = result.financial_metrics
    key_metrics = [
        ("3-Year ROI", f"{fm.three_year_roi_percentage:,.1f}%"),
        ("Payback Period", f"{fm.payback_period_months:.1f} months"),
        ("3-Year Net Present Value", f"QAR {fm.npv_3yr_qar:,.0f}"),
        ("Benefit-Cost Ratio", f"{fm.benefit_cost_ratio:.2f}x"),
        ("Annual Net Benefit", f"QAR {fm.annual_net_benefit_qar:,.0f}"),
        ("Confidence Level", result.confidence_level),
    ]
    _kv_table(doc, key_metrics)

    doc.add_page_break()

    # ── 2. USE CASE OVERVIEW ──────────────────────────────────
    _heading(doc, "2. Use Case Overview")
    _add_horizontal_rule(doc)

    _heading(doc, "2.1 Use Case Description", level=2)
    _body(doc, inp.description)

    _heading(doc, "2.2 Current Process", level=2)
    _body(doc, inp.current_process)

    _heading(doc, "2.3 Pain Points & Challenges", level=2)
    _body(doc, inp.pain_points)

    _heading(doc, "2.4 Proposed AI Solution", level=2)
    ai_types_text = ", ".join(inp.ai_types)
    _body(doc, f"AI Solution Types: {ai_types_text}")
    _body(doc, f"Primary Benefit Target: {inp.primary_benefit}")
    _body(doc, f"Expected Beneficiaries: {inp.beneficiary_count:,} users/citizens")

    doc.add_page_break()

    # ── 3. ROI ANALYSIS ───────────────────────────────────────
    _heading(doc, "3. ROI Analysis")
    _add_horizontal_rule(doc)

    _heading(doc, "3.1 Methodology & Assumptions", level=2)
    assumptions = [
        ("Analysis Horizon", "3 Years"),
        ("Discount Rate", "8% (Qatar government benchmark)"),
        ("FTEs Involved", str(inp.fte_count)),
        ("Hours/Week per FTE on Task", f"{inp.hours_per_week_per_fte}h"),
        ("Average Annual Salary", inp.salary_range),
        ("Current Error Rate", inp.error_rate),
        ("Data Availability", inp.data_availability),
        ("Estimated Automation Rate", f"{result.automation_percentage * 100:.0f}%"),
        ("Time Reduction Rate", f"{result.time_reduction_percentage * 100:.0f}%"),
        ("Error Reduction Rate", f"{result.error_reduction_percentage * 100:.0f}%"),
    ]
    _kv_table(doc, assumptions)

    _heading(doc, "3.2 Annual Benefits Breakdown", level=2)
    _benefits_table(doc, result)

    _heading(doc, "3.3 Financial Metrics", level=2)
    _financial_table(doc, result)

    _heading(doc, "3.4 Implementation Cost Estimate", level=2)
    cost_rows = [
        ("Conservative Estimate", f"QAR {result.implementation_cost_low_qar:,.0f}"),
        ("Mid-Range Estimate", f"QAR {result.implementation_cost_mid_qar:,.0f}"),
        ("High Estimate", f"QAR {result.implementation_cost_high_qar:,.0f}"),
        ("Annual Maintenance Cost", f"QAR {result.annual_maintenance_cost_qar:,.0f}"),
        ("Preferred Implementation Timeline", f"{inp.preferred_timeline_months} months"),
        ("Budget Range Indicated", inp.budget_range),
    ]
    _kv_table(doc, cost_rows)

    doc.add_page_break()

    # ── 4. IMPACT ASSESSMENT ──────────────────────────────────
    _heading(doc, "4. Impact Assessment")
    _add_horizontal_rule(doc)

    _heading(doc, "4.1 Operational Impact", level=2)
    _body(doc, result.operational_impact)

    _heading(doc, "4.2 Strategic Impact", level=2)
    _body(doc, result.strategic_impact)

    _heading(doc, "4.3 Citizen & Employee Impact", level=2)
    _body(doc, result.citizen_impact)

    _heading(doc, "4.4 Qatar Vision 2030 Alignment", level=2)
    _body(doc, result.vision_2030_alignment)

    doc.add_page_break()

    # ── 5. RISK ASSESSMENT ────────────────────────────────────
    _heading(doc, "5. Risk Assessment")
    _add_horizontal_rule(doc)
    _risk_table(doc, result.risks)

    doc.add_page_break()

    # ── 6. IMPLEMENTATION ROADMAP ─────────────────────────────
    _heading(doc, "6. Implementation Roadmap")
    _add_horizontal_rule(doc)

    for phase_data in result.implementation_phases:
        _heading(doc, phase_data.get("phase", "Phase"), level=2)
        duration_para = doc.add_paragraph(f"Duration: {phase_data.get('duration', 'TBD')}")
        duration_para.runs[0].italic = True
        duration_para.runs[0].font.size = Pt(10)

        activities = phase_data.get("activities", [])
        if activities:
            act_para = doc.add_paragraph()
            act_run = act_para.add_run("Key Activities:")
            act_run.bold = True
            act_run.font.size = Pt(10)
            for act in activities:
                bullet = doc.add_paragraph(f"• {act}", style="Normal")
                bullet.paragraph_format.left_indent = Inches(0.3)
                bullet.runs[0].font.size = Pt(10)

        deliverables = phase_data.get("deliverables", [])
        if deliverables:
            del_para = doc.add_paragraph()
            del_run = del_para.add_run("Deliverables:")
            del_run.bold = True
            del_run.font.size = Pt(10)
            for d in deliverables:
                bullet = doc.add_paragraph(f"✓ {d}", style="Normal")
                bullet.paragraph_format.left_indent = Inches(0.3)
                bullet.runs[0].font.size = Pt(10)
                bullet.runs[0].font.color.rgb = RGBColor(0x15, 0x6A, 0x38)

    doc.add_page_break()

    # ── 7. RECOMMENDATIONS ────────────────────────────────────
    _heading(doc, "7. Recommendations")
    _add_horizontal_rule(doc)
    _body(doc, result.recommendations)

    doc.add_paragraph()
    _heading(doc, "Strategic Priority Level", level=2)
    priority_para = doc.add_paragraph()
    priority_run = priority_para.add_run(f"  {inp.strategic_priority.upper()}  ")
    priority_run.bold = True
    priority_run.font.size = Pt(14)
    priority_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    priority_colors = {
        "Critical": "8B172E",
        "High": "C0392B",
        "Medium": "E67E22",
        "Low": "27AE60",
    }
    # Note: inline background color requires complex XML; we use bold colored text instead
    priority_run.font.color.rgb = RGBColor(
        *bytes.fromhex(priority_colors.get(inp.strategic_priority, "0D1B3E"))
    )

    doc.add_page_break()

    # ── 8. APPENDIX ───────────────────────────────────────────
    _heading(doc, "8. Appendix – Submission Details")
    _add_horizontal_rule(doc)
    appendix_rows = [
        ("Submitter Name", inp.submitter_name),
        ("Submitter Email", inp.submitter_email),
        ("Sector", inp.sector),
        ("Department", inp.department),
        ("Data Types Available", ", ".join(inp.data_types) if inp.data_types else "Not specified"),
        ("Preferred Timeline", f"{inp.preferred_timeline_months} months"),
        ("Report Generated", datetime.now().strftime("%d %B %Y, %H:%M")),
    ]
    _kv_table(doc, appendix_rows)

    _add_horizontal_rule(doc)
    footer_para = doc.add_paragraph(
        "This report is generated by the NPC AI ROI Evaluation System. "
        "All financial estimates are indicative and should be validated by the Finance and Procurement departments "
        "before formal commitment. © National Planning Council, State of Qatar."
    )
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    footer_para.runs[0].italic = True

    # Serialize to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
